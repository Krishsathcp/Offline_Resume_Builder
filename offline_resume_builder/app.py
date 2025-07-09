from flask import Flask, render_template, request, send_file, jsonify, send_from_directory, redirect, url_for
from offline_resume_builder.utils.pdf_generator import generate_pdf_from_html
from offline_resume_builder.utils.ollama_ai import (
    generate_ai_resume_content,
    convert_text_to_education_list,
    convert_text_to_experience_list,
    generate_ai_resume_sections
)
from offline_resume_builder.utils.ai_engine import query_mistral as query_llama3
from docx import Document
from docx.shared import Inches
from datetime import datetime
from bs4 import BeautifulSoup
import uuid
import os
import json
import random
import ollama
import pdfkit
import re
import zipfile
import mammoth
import shutil
base_dir = os.path.dirname(os.path.abspath(__file__))
history_dir = os.path.join(base_dir, "..", "history")

#if os.path.exists(history_dir):
 #   try:
    #    shutil.rmtree(history_dir)
     #   print("✅ Deleted old history directory")
   # except Exception as e:
    #    print("❌ Could not delete history:", e)

#os.makedirs(history_dir, exist_ok=True)
#print("✅ Recreated clean writable history directory")

def strip_tags(html):
    """Remove HTML tags from text"""
    return re.sub('<[^<]+?>', '', html)

WKHTMLTOPDF_PATH = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe'

app = Flask(__name__)
app.jinja_env.globals['now'] = datetime.now

@app.route('/')
def home():
    return render_template('form.html')

def html_to_plaintext_bullets(html):
    """Convert HTML lists to plaintext bullet points"""
    soup = BeautifulSoup(html, 'html.parser')
    lines = []
    
    for ul in soup.find_all(['ul', 'ol']):
        for li in ul.find_all('li', recursive=False):
            lines.append(f"- {li.get_text(strip=True)}")
    
    if not lines:
        return soup.get_text(separator="\n").strip()
    
    return "\n".join(lines)

@app.route('/generate_summary', methods=['POST'])
def generate_summary():
    data = request.get_json()
    skills = data.get("skills", "")
    experience = data.get("experience", "")
    
    prompt = f"""Write a concise 3-4 line professional summary for a resume based on the following:

Skills:
{skills}

Experience:
{experience}
"""
    
    try:
        res = query_llama3(prompt)
        print("AI summary response >>>", res)
        return jsonify({"summary": res.strip()})
    except Exception as e:
        return jsonify({"summary": "AI failed: " + str(e)})

@app.route('/edit_resume/<resume_id>')
def edit_resume(resume_id):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    history_file = os.path.join(base_dir, "..", "history", "history.json")

    if os.path.exists(history_file):
        with open(history_file, 'r', encoding='utf-8') as f:
            try:
                history_data = json.load(f)
                for item in history_data:
                    if item.get("id") == resume_id:
                        return render_template("form.html", resume=item, editing=True)
            except json.JSONDecodeError:
                pass

    return "Resume not found", 404


@app.route('/preview_docx/<resume_id>')
def preview_docx(resume_id):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    history_file = os.path.join(base_dir, "..", "history", "history.json")

    if not os.path.exists(history_file):
        return "No resume history found.", 404

    try:
        with open(history_file, 'r', encoding='utf-8') as f:
            history = json.load(f)
    except json.JSONDecodeError:
        return "Resume history is corrupted.", 500

    entry = next((item for item in history if item['id'] == resume_id), None)
    if not entry:
        return "Resume entry not found.", 404

    template_style = entry.get("template", "formal").lower()
    if template_style not in ["formal", "casual"]:
        template_style = "formal"

    return render_template(
        f"{template_style}_resume_template.html",
        name=entry.get("name", ""),
        email=entry.get("email", ""),
        phone=entry.get("phone", ""),
        address=entry.get("address", ""),
        ai_summary=entry.get("summary", ""),
        ai_suggestions=entry.get("ai_suggestions", ""),
        skills_html=entry.get("skills", "").replace("\n", "<br>"),
        experience_html=entry.get("experience", "").replace("\n", "<br>"),
        education_html=entry.get("education", "").replace("\n", "<br>"),
        projects_html=entry.get("projects", "").replace("\n", "<br>"),
        extra_sections=entry.get("extra_sections", []),
        user_suggestions=entry.get("user_suggestions", "")
    )


@app.route('/export_all_resumes')
def export_all_resumes():
    history_file = 'history/history.json'
    zip_filename = "all_resumes.zip"
    zip_path = os.path.join("temp", zip_filename)
    
    os.makedirs("temp", exist_ok=True)
    
    files_to_export = []
    
    if os.path.exists(history_file):
        with open(history_file, 'r', encoding='utf-8') as f:
            try:
                history_data = json.load(f)
                for item in history_data:
                    file_path = item.get("file_path")
                    if file_path and os.path.exists(file_path):
                        files_to_export.append(file_path)
            except json.JSONDecodeError:
                return "Failed to load resume history.", 500
    
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for file_path in files_to_export:
            zipf.write(file_path, arcname=os.path.basename(file_path))
    
    return send_file(zip_path, as_attachment=True)

@app.route('/improve_resume', methods=['GET', 'POST'])
def improve_resume():
    suggestion = ""
    resume_text = ""
    job_description = ""
    
    if request.method == 'POST':
        resume_text = request.form.get('resume_text', '')
        job_description = request.form.get('job_description', '')
        
        if resume_text.strip():
            prompt = f"Suggest improvements for this resume:\n\n{resume_text}"
            suggestion = query_llama3(prompt)
        else:
            suggestion = "Please paste your resume text first."
    
    return render_template(
        'resume_improvement.html',
        suggestions=suggestion,
        resume_text=resume_text,
        job_description=job_description
    )

interview_qbank = []

@app.route('/temp/<filename>')
def serve_temp_file(filename):
    return send_from_directory('temp', filename)

@app.route('/score_resume', methods=['GET', 'POST'])
def score_resume():
    score = None
    feedback = ""
    if request.method == 'POST':
        resume_text = request.form.get('resume_text', '')
        if resume_text:
            prompt = f"""
You are a professional resume reviewer.

Evaluate the following resume and return:
1. A score out of 100 based on formatting, relevance, clarity, and impact.
2. Detailed feedback to improve the resume.

Resume:
\"\"\"
{resume_text}
\"\"\"
"""
            try:
                response = query_llama3(prompt)
                score_match = re.search(r'(\d{1,3})/100', response)
                score = score_match.group(1) if score_match else None
                feedback = response
            except Exception as e:
                feedback = f"AI error: {e}"
    
    return render_template('score_resume.html', score=score, feedback=feedback)

@app.route('/resumes/<filename>')
def download_resume(filename):
    import json
    from offline_resume_builder.utils.resume_templates import render_resume_to_pdf, render_resume_to_docx

    resumes_folder = "resumes"
    history_path = os.path.join("history", "history.json")
    filepath = os.path.join(resumes_folder, filename)

    if os.path.exists(filepath):
        return send_from_directory(resumes_folder, filename, as_attachment=True)

    resume_id = filename.rsplit('.', 1)[0]
    extension = filename.rsplit('.', 1)[-1].lower()

    try:
        with open(history_path, "r", encoding="utf-8") as f:
            history = json.load(f)
    except Exception as e:
        print("⚠️ Failed to read history:", e)
        return "History not found", 500

    resume_data = next((r for r in history if r["id"] == resume_id), None)
    if not resume_data:
        print("❌ Resume ID not found in history:", resume_id)
        return "Resume data not found", 404

    try:
        if extension == "pdf":
            pdf_bytes = render_resume_to_pdf(resume_data)
            with open(filepath, "wb") as f:
                f.write(pdf_bytes)

        elif extension == "docx":
            docx_bytes = render_resume_to_docx(resume_data)
            with open(filepath, "wb") as f:
                f.write(docx_bytes)
        else:
            return "Unsupported file format", 400
    except Exception as e:
        print("❌ Failed to regenerate resume:", e)
        return "Resume generation failed", 500

    return send_from_directory(resumes_folder, filename, as_attachment=True)

@app.route('/interview_prep', methods=['GET', 'POST'])
def interview_prep():
    global interview_qbank
    feedback = ""
    current_questions = []
    answers = []
    
    if not interview_qbank:
        prompt = "Give me 15 short, tough and varied job interview questions for a software developer. Number them."
        response = query_llama3(prompt)
        interview_qbank = [q.strip() for q in response.split("\n") if q.strip()]
        random.shuffle(interview_qbank)
    
    if request.method == 'POST':
        answers = [request.form.get(f'answer{i}', '') for i in range(3)]
        qna = "\n".join([f"Q{i+1}: {request.form.get(f'q{i}')}\nA{i+1}: {answers[i]}" for i in range(3)])
        feedback_prompt = f"Score each answer out of 10 and provide feedback.\n\n{qna}"
        feedback = query_llama3(feedback_prompt)
    
    current_questions = [interview_qbank.pop() for _ in range(min(3, len(interview_qbank)))]
    
    return render_template('interview_prep.html', questions=current_questions, feedback=feedback, answers=answers)

@app.route('/get_mcqs', methods=['POST'])
def get_mcqs():
    topic = request.json.get('topic', '')
    if not topic:
        return jsonify({'error': 'No topic provided'}), 400
    
    prompt = f"Generate 3 easy MCQ interview questions for the topic '{topic}'. Each question should have 4 options and mark the correct one clearly. Format:\nQ: ...\nA) ...\nB) ...\nC) ...\nD) ...\nAnswer: ..."
    
    response = query_llama3(prompt)
    questions = []
    
    if response:
        blocks = response.strip().split("Q:")
        for block in blocks[1:]:
            lines = block.strip().split("\n")
            if not lines or len(lines) < 5:
                continue
            
            q_text = lines[0].strip()
            
            option_pattern = re.compile(r'^[A-D][).:-]?\s+(.*)$')
            options = []
            for line in lines[1:]:
                match = option_pattern.match(line.strip())
                if match:
                    options.append(match.group(1))
            
            answer_line = next((line for line in lines if line.lower().startswith("answer:")), "")
            answer = answer_line.split(":")[1].strip() if ":" in answer_line else ""
            
            if len(options) == 4:
                questions.append({
                    'question': q_text,
                    'options': options,
                    'answer': answer
                })
        
        return jsonify({'questions': questions})
    else:
        return jsonify({'error': 'Failed to generate questions'}), 500

@app.route('/resume_improvement', methods=['GET', 'POST'])
def resume_improvement():
    suggestions = None
    
    if request.method == 'POST':
        resume_text = request.form['resume_text']
        job_description = request.form.get('job_description', '')
        
        prompt = f"""
            You are an expert resume advisor.

            Here is the resume:
            \"\"\"
            {resume_text}
            \"\"\"

            Here is the job description (if provided):
            \"\"\"
            {job_description}
            \"\"\"

            Give me the following:
            1. Resume improvement suggestions (grammar, clarity, formatting, etc.)
            2. How well this resume matches the job description
            3. Tips to better align the resume with the job
            4. Use markdown or bullet points if needed
        """
        
        try:
            response = ollama.chat(model='mistral', messages=[
                {"role": "user", "content": prompt}
            ])
            suggestions = response['message']['content']
        except Exception as e:
            suggestions = f"Error communicating with AI model: {e}"
    
    return render_template('resume_improvement.html', suggestions=suggestions)

@app.route('/download_improved_resume', methods=['POST'])
def download_improved_resume():
    improved_text = request.form['improved_text']
    template_choice = request.form['template_choice']
    
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    temp_dir = os.path.join(base_dir, "temp")
    resumes_dir = os.path.join(base_dir, "resumes")
    
    os.makedirs(temp_dir, exist_ok=True)
    os.makedirs(resumes_dir, exist_ok=True)
    
    filename = f"improved_{uuid.uuid4().hex}"
    html_path = os.path.join(temp_dir, f"{filename}.html")
    pdf_path = os.path.join(resumes_dir, f"{filename}.pdf")
    
    html_content = f"""
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{
            font-family: Arial, sans-serif;
            margin: 40px;
            line-height: 1.6;
            color: #333;
        }}
        h2 {{
            color: #0052cc;
        }}
        ul {{
            margin-top: 10px;
        }}
        li {{
            margin-bottom: 6px;
        }}
    </style>
</head>
<body>
    <h2>AI Suggestions</h2>
    <div>{improved_text.replace(chr(10), '<br>')}</div>
</body>
</html>
"""
    
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    
    config = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF_PATH)
    
    try:
        pdfkit.from_file(html_path, pdf_path, configuration=config)
    except Exception as e:
        return f"❌ PDF generation failed: {e}", 500
    
    if not os.path.exists(pdf_path):
        return "❌ PDF file not generated properly.", 500
    
    return send_file(pdf_path, as_attachment=True)

@app.route('/history')
def view_history():
    sort_order = request.args.get('sort_order', 'desc')
    file_type = request.args.get('file_type', 'all')    

    base_dir = os.path.dirname(os.path.abspath(__file__))
    history_file = os.path.join(base_dir, "..", "history", "history.json")
    history_data = []

    if os.path.exists(history_file):
        with open(history_file, 'r', encoding='utf-8') as f:
            try:
                history_data = json.load(f)
            except json.JSONDecodeError:
                history_data = []

    if file_type in ['pdf', 'docx']:
        history_data = [item for item in history_data if item.get('format') == file_type]

    history_data.sort(
        key=lambda x: x.get('date', ''),
        reverse=(sort_order == 'desc')
    )

    return render_template(
        "resume_history.html",
        history=history_data,
        sort_order=sort_order,
        file_type=file_type
    )



@app.route('/preview_resume/<resume_id>')
def preview_resume(resume_id):
    base_dir = os.path.dirname(os.path.abspath(__file__))
    history_file = os.path.join(base_dir, "..", "history", "history.json")

    if not os.path.exists(history_file):
        return "Resume history file not found.", 404

    try:
        with open(history_file, "r", encoding="utf-8") as f:
            history = json.load(f)
    except json.JSONDecodeError:
        return "History file is corrupted.", 500

    entry = next((item for item in history if item["id"] == resume_id), None)
    if not entry:
        return "Resume entry not found in history.", 404

    template_style = entry.get("template", "formal").lower()
    if template_style not in ["formal", "casual"]:
        template_style = "formal"

    return render_template(
        f"{template_style}_resume_template.html",
        name=entry.get("name", ""),
        email=entry.get("email", ""),
        phone=entry.get("phone", ""),
        address=entry.get("address", ""),
        ai_summary=entry.get("summary", ""),
        ai_suggestions=entry.get("ai_suggestions", ""),
        skills_html=entry.get("skills", "").replace("\n", "<br>"),
        experience_html=entry.get("experience", "").replace("\n", "<br>"),
        education_html=entry.get("education", "").replace("\n", "<br>"),
        projects_html=entry.get("projects", "").replace("\n", "<br>"),
        extra_sections=entry.get("extra_sections", []),
        user_suggestions=entry.get("user_suggestions", "")
    )



@app.route('/delete_resume', methods=['POST'])
def delete_resume():
    resume_id = request.form.get("resume_id")

    if not resume_id:
        return "Missing resume ID", 400

    base_dir = os.path.dirname(os.path.abspath(__file__))
    history_dir = os.path.join(base_dir, "..", "history")
    history_file = os.path.join(history_dir, "history.json")

    if os.path.exists(history_file):
        with open(history_file, 'r', encoding='utf-8') as f:
            try:
                history_data = json.load(f)
            except json.JSONDecodeError:
                history_data = []

        updated_history = []
        for item in history_data:
            if item["id"] == resume_id:
                try:
                    file_path = item.get("file_path")
                    if file_path and os.path.exists(file_path):
                        os.remove(file_path)
                    temp_html = os.path.join("temp", f"{resume_id}.html")
                    if os.path.exists(temp_html):
                        os.remove(temp_html)
                except Exception as e:
                    print(f"Failed to delete associated files: {e}")
            else:
                updated_history.append(item)

        with open(history_file, 'w', encoding='utf-8') as f:
            json.dump(updated_history, f, indent=2)

    return redirect(url_for('view_history'))


@app.route('/summarize_ai_full', methods=['POST'])
def summarize_ai_full():
    data = request.get_json()

    summary = data.get("summary", "").strip()
    skills = data.get("skills", "").strip()
    experience = data.get("experience", "").strip()
    education = data.get("education", "").strip()
    projects = data.get("projects", "").strip()
    job_description = data.get("job_description", "").strip()
    extra_sections = data.get("extra_sections", [])

    if not job_description:
        return jsonify({"error": "Job description is required for AI enhancement."}), 400

    extra_text = ""
    if extra_sections:
        for section in extra_sections:
            heading = section.get("heading", "").strip()
            content = section.get("content", "").strip()
            if heading:
                extra_text += f"\n{heading}:\n{content or 'N/A'}"

    prompt = f"""
You are an expert resume optimizer.

Below is the candidate's resume:

SUMMARY:
{summary or "N/A"}

SKILLS:
{skills or "N/A"}

EXPERIENCE:
{experience or "N/A"}

EDUCATION:
{education or "N/A"}

PROJECTS:
{projects or "N/A"}

EXTRA SECTIONS:
{extra_text or "None"}

JOB DESCRIPTION:
{job_description}

TASK:
- Improve each section INCLUDING only the provided EXTRA SECTIONS.
- ❗ DO NOT invent or rename extra sections. Keep headings exactly as given (case/style can vary).
- Return improvements in this exact structure:

SUMMARY:
...

SKILLS:
...

EXPERIENCE:
...

EDUCATION:
...

PROJECTS:
...

EXTRA SECTIONS:
- Heading1:
  Content...
- Heading2:
  Content...

CHANGES MADE:
- ...
"""


    try:
        response = query_llama3(prompt)
        print("AI Response >>>", response)

        def extract_block(label, text):
            start = text.find(f"{label.upper()}:")
            if start == -1:
                return ""
            labels = ["SUMMARY", "SKILLS", "EXPERIENCE", "EDUCATION", "PROJECTS", "EXTRA SECTIONS", "CHANGES MADE"]
            next_starts = [text.find(f"{l}:", start + 1) for l in labels if l != label.upper()]
            next_starts = [p for p in next_starts if p != -1]
            end = min(next_starts) if next_starts else None
            return text[start + len(label) + 1:end].strip() if end else text[start + len(label) + 1:].strip()

        ai_summary = extract_block("SUMMARY", response)
        ai_skills = extract_block("SKILLS", response)
        ai_experience = extract_block("EXPERIENCE", response)
        ai_education = extract_block("EDUCATION", response)
        ai_projects = extract_block("PROJECTS", response)
        ai_changes = extract_block("CHANGES MADE", response)

        ai_extra_sections_raw = extract_block("EXTRA SECTIONS", response)
        structured_extra = []
        current_heading = ""
        if extra_sections:  
            lines = ai_extra_sections_raw.splitlines()
            current_content = []

            for line in lines:
                line = line.strip()
                if line.startswith("- ") and ":" in line:
                    if current_heading:
                        structured_extra.append({
                            "heading": current_heading,
                            "content": "\n".join(current_content).strip()
                        })
                        current_content = []

                    heading_line = line[2:].strip()
                    heading_parts = heading_line.split(":", 1)
                    raw_heading = heading_parts[0].strip()
                    formatted_heading = raw_heading.capitalize()
                    rest = heading_parts[1].strip() if len(heading_parts) > 1 else ""

                    current_heading = formatted_heading
                    if rest:
                        current_content = [rest]
                else:
                    current_content.append(line)

            if current_heading:
                structured_extra.append({
                    "heading": current_heading,
                    "content": "\n".join(current_content).strip()
                })

        if current_heading:
            structured_extra.append({
                "heading": current_heading,
                "content": "\n".join(current_content).strip()
            })

        original_extra_headings = {
            section.get("heading", "").strip().lower()
            for section in extra_sections if section.get("heading", "").strip()
        }

        filtered_extra = [
            {
                "heading": sec["heading"],
                "content": sec["content"]
            }
            for sec in structured_extra
            if sec["heading"].strip().lower() in original_extra_headings
        ]
              
        discarded = [
                sec["heading"]
                for sec in structured_extra
                if sec["heading"].strip().lower() not in original_extra_headings
        ]
        if discarded:
            print("❌ Discarded invented AI sections:", discarded)
        
        return jsonify({
            "ai_summary": ai_summary,
            "skills": ai_skills,
            "experience": ai_experience,
            "education": ai_education,
            "projects": ai_projects,
            "extra_sections": filtered_extra,  # ✅ Cleaned
            "changes": ai_changes
        })



    except Exception as e:
        return jsonify({"error": str(e)}), 500









@app.route('/generate_resume', methods=['POST'])
def generate_resume():
    data = request.form

    name = data.get("name")
    email = data.get("email")
    phone = data.get("phone")
    address = data.get("address")
    summary = data.get("summary", "")
    job_description = data.get("job_description", "")
    template_style = data.get("style", "formal").lower()
    export_format = data.get("export_format", "pdf").lower()
    use_ai = data.get("use_ai_final", "false") == "true"

    raw_skills = data.get("skills", "")
    raw_experience = data.get("experience", "")
    raw_education = data.get("education", "")
    raw_projects = data.get("projects", "")
    user_suggestions = data.get("user_suggestions", "")
    extra_sections_raw = data.get("extra_sections", "[]")

    try:
        extra_sections = json.loads(extra_sections_raw)
    except json.JSONDecodeError:
        extra_sections = []

    if template_style not in ["formal", "casual"]:
        template_style = "formal"

    ai_sections = generate_ai_resume_sections(
        raw_skills, raw_experience, raw_education, raw_projects, job_description
    )

    ai_summary_block = generate_ai_resume_content({
        "name": name,
        "summary": summary,
        "skills": raw_skills.split(","),
        "experience": raw_experience,
        "education": raw_education,
        "projects": raw_projects,
        "user_suggestions": user_suggestions,
        "extra_sections": extra_sections,
    })

    output_id = uuid.uuid4().hex
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    temp_dir = os.path.join(base_dir, "temp")
    resumes_dir = os.path.join(base_dir, "resumes")
    os.makedirs(temp_dir, exist_ok=True)
    os.makedirs(resumes_dir, exist_ok=True)

    content = {
        "summary": ai_summary_block.get("ai_summary", summary) if use_ai else summary,
        "suggestions": ai_summary_block.get("suggestions", "") if use_ai else "",
        "skills": ai_sections["skills"] if use_ai else raw_skills.replace("\n", "<br>"),
        "experience": ai_sections["experience"] if use_ai else raw_experience.replace("\n", "<br>"),
        "education": ai_sections["education"] if use_ai else raw_education.replace("\n", "<br>"),
        "projects": ai_sections["projects"] if use_ai else raw_projects.replace("\n", "<br>"),
    }
   
    html_content = render_template(
        f"{template_style}_resume_template.html",
        name=name,
        email=email,
        phone=phone,
        address=address,
        ai_summary=content["summary"],
        ai_suggestions=content["suggestions"],
        skills_html=content["skills"],
        experience_html=content["experience"],
        education_html=content["education"],
        projects_html=content["projects"],
        extra_sections=extra_sections,
        user_suggestions=user_suggestions
    )

    html_path = os.path.join(temp_dir, f"{output_id}.html")
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)

    history_data = {
        "id": output_id,
        "name": name,
        "email": email,
        "phone": phone,
        "address": address,
        "summary": content["summary"],
        "job_description": job_description,
        "template": template_style,
        "skills": raw_skills,
        "experience": raw_experience,
        "education": raw_education,
        "projects": raw_projects,
        "ai_suggestions": content["suggestions"],
        "user_suggestions": user_suggestions,
        "extra_sections": extra_sections,
        "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }

    if export_format == "docx":
        from docx import Document
        doc = Document()
        doc.add_heading(name, 0)
        doc.add_paragraph(f"Email: {email} | Phone: {phone} | Address: {address}")

        if content["summary"]:
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(content["summary"])
        if content["education"]:
            doc.add_heading("Education", level=1)
            doc.add_paragraph(strip_tags(content["education"]))
        if content["experience"]:
            doc.add_heading("Experience", level=1)
            doc.add_paragraph(strip_tags(content["experience"]))
        if content["projects"]:
            doc.add_heading("Projects", level=1)
            doc.add_paragraph(strip_tags(content["projects"]))
        if content["skills"]:
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(strip_tags(content["skills"]))
        if content["suggestions"]:
            doc.add_heading("AI Suggestions", level=1)
            doc.add_paragraph(content["suggestions"])
        if user_suggestions:
            doc.add_heading("User Suggestions", level=1)
            doc.add_paragraph(strip_tags(user_suggestions))

        for section in extra_sections:
            if section.get("heading") and section.get("content"):
                doc.add_heading(section["heading"], level=1)
                doc.add_paragraph(strip_tags(section["content"]))

        docx_path = os.path.join(resumes_dir, f"{output_id}.docx")
        doc.save(docx_path)

        with open(os.path.join(temp_dir, f"{output_id}.html"), 'w', encoding='utf-8') as f:
            f.write(html_content)

        history_data["format"] = "docx"
        history_data["file_path"] = f"resumes/{output_id}.docx"

        _save_history(history_data)
        return send_file(docx_path, as_attachment=True)

    
    pdf_path = os.path.join(resumes_dir, f"{output_id}.pdf")
    generate_pdf_from_html(html_path, pdf_path)

    history_data["format"] = "pdf"
    history_data["file_path"] = f"resumes/{output_id}.pdf"

    _save_history(history_data)
    return send_file(pdf_path, as_attachment=True)



def extract_missing_skills(user_skills, job_description):
    """Extract skills missing from user's resume based on job description"""
    prompt = f"""
You are an AI resume assistant.

Candidate's skills: {user_skills}
Job description:
\"\"\"
{job_description}
\"\"\"

List 5 to 10 important skills required in this job that are missing from the candidate's list. 
Return only the skill names, comma-separated.
"""
    try:
        response = query_llama3(prompt)
        return [skill.strip() for skill in response.split(",") if skill.strip()]
    except Exception as e:
        print("Skill extraction error:", e)
        return []

def _save_history(new_entry):
    """Helper function to save resume history"""
    import os, json

    base_dir = os.path.dirname(os.path.abspath(__file__))
    history_dir = os.path.join(base_dir, "..", "history")
    os.makedirs(history_dir, exist_ok=True)
    history_file = os.path.join(history_dir, "history.json")

    try:
        if os.path.exists(history_file):
            with open(history_file, "r", encoding="utf-8") as f:
                existing = json.load(f)
        else:
            existing = []

        existing.append(new_entry)
        with open(history_file, "w", encoding="utf-8") as f:
            json.dump(existing, f, indent=2)
    except Exception as e:
        print("❌ Failed to save resume history:", e)


def main():
    app.run(debug=True)

if __name__ == "__main__":
    main()