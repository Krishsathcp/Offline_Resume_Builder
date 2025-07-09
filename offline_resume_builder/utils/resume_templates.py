import pdfkit
from docx import Document
from jinja2 import Environment, FileSystemLoader
import os

env = Environment(loader=FileSystemLoader('offline_resume_builder/templates'))

def render_resume_to_pdf(resume):
    template = env.get_template(f"{resume['template']}_resume_template.html")
    html = template.render(resume=resume)
    return pdfkit.from_string(html, False)

def render_resume_to_docx(resume):
    template = env.get_template(f"{resume['template']}_resume_template.html")
    html = template.render(resume=resume)

    doc = Document()
    doc.add_paragraph("Generated Resume")
    doc.add_paragraph(" ")

    import html2text
    text = html2text.html2text(html)
    doc.add_paragraph(text)

    from io import BytesIO
    f = BytesIO()
    doc.save(f)
    return f.getvalue()
